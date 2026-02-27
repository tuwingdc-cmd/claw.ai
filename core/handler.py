"""
Message Handler + DB-backed Conversation Memory + Smart Skills + Tools + Music + URL Fetch
"""
import json
import math
import logging
import re
import asyncio
import aiohttp
import os
import tempfile
import shutil
from typing import Dict, List, Optional
from datetime import datetime, timedelta
from core.providers import ProviderFactory, AIResponse
from core.database import (
    save_message, get_conversation, clear_conversation,
    get_memory_stats, MAX_MEMORY_MESSAGES
)
from config import API_KEYS, FALLBACK_CHAINS, PROVIDERS

log = logging.getLogger(__name__)

# Re-export for backward compatibility
MEMORY_EXPIRE_MINUTES = 0

# ============================================================
# REQUEST LOGS
# ============================================================

request_logs: List[Dict] = []
MAX_LOGS = 500

def _log_request(guild_id, provider, model, success, latency, is_fallback=False, error=None):
    request_logs.append({"guild_id": guild_id, "provider": provider, "model": model, "success": success, "latency": latency, "is_fallback": is_fallback, "error": error, "time": datetime.now().strftime("%H:%M:%S")})
    if len(request_logs) > MAX_LOGS:
        request_logs.pop(0)

def strip_think_tags(content: str) -> str:
    for tag in ['think', 'thinking', 'thought']:
        content = re.sub(rf'<{tag}>.*?</{tag}>', '', content, flags=re.DOTALL)
        content = re.sub(rf'</?{tag}>', '', content)
    return re.sub(r'\n{3,}', '\n\n', content).strip()

# ============================================================
# SEARCH — Tavily first, DuckDuckGo fallback
# ============================================================

async def do_search(query: str, engine: str = "auto") -> str:
    tavily_key = API_KEYS.get("tavily")
    if tavily_key:
        try:
            async with aiohttp.ClientSession() as session:
                async with session.post(
                    "https://api.tavily.com/search",
                    json={
                        "api_key": tavily_key,
                        "query": query,
                        "max_results": 5,
                        "search_depth": "basic",
                        "include_answer": True,
                    },
                    timeout=aiohttp.ClientTimeout(total=15)
                ) as resp:
                    if resp.status == 200:
                        data = await resp.json()
                        parts = []
                        if data.get("answer"):
                            parts.append(f"Summary: {data['answer']}")
                        for i, r in enumerate(data.get("results", []), 1):
                            parts.append(
                                f"{i}. {r.get('title', 'No title')}\n"
                                f"   {r.get('content', '')[:200]}\n"
                                f"   {r.get('url', '')}"
                            )
                        if parts:
                            log.info(f"🔍 Tavily search OK: {query}")
                            return "\n\n".join(parts)
                    else:
                        log.warning(f"Tavily HTTP {resp.status}, fallback to DuckDuckGo")
        except Exception as e:
            log.warning(f"Tavily error, fallback to DuckDuckGo: {e}")

    try:
        from duckduckgo_search import DDGS
        def _s():
            with DDGS() as d:
                return list(d.text(query, max_results=5))
        results = await asyncio.get_event_loop().run_in_executor(None, _s)
        if not results:
            return "Tidak ada hasil."
        log.info(f"🔍 DuckDuckGo search OK: {query}")
        return "\n\n".join([
            f"{i}. {r['title']}\n   {r['body'][:150]}\n   {r['href']}"
            for i, r in enumerate(results, 1)
        ])
    except Exception as e:
        return f"Search error: {e}"

GROUNDING_MODELS = {("groq", "groq/compound"), ("groq", "groq/compound-mini"), ("groq", "compound-beta"), ("groq", "compound-beta-mini"), ("pollinations", "gemini-search"), ("pollinations", "perplexity-fast"), ("pollinations", "perplexity-reasoning")}
def is_grounding_model(p, m): return (p, m) in GROUNDING_MODELS

# ============================================================
# TRANSLATE — AI-powered natural translation
# ============================================================

async def do_translate(text: str, target_lang: str, style: str = "natural") -> str:
    style_prompts = {
        "natural": (
            "You are a native bilingual translator. Translate naturally as a native speaker would say it. "
            "Understand slang, idioms, colloquialisms, and cultural context. "
            "Do NOT translate literally word-by-word. Make it sound like something a real person would say. "
            "Only output the translation, nothing else."
        ),
        "formal": (
            "You are a professional translator. Translate in formal, polished language. "
            "Only output the translation, nothing else."
        ),
        "casual": (
            "You are a young native speaker. Translate super casually with modern slang and abbreviations. "
            "Make it sound like texting a friend. Only output the translation, nothing else."
        ),
    }

    system_prompt = style_prompts.get(style, style_prompts["natural"])
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": f"Translate to {target_lang}:\n\n{text}"}
    ]

    translate_chains = [
        ("groq", "llama-3.1-8b-instant"),
        ("groq", "llama-3.3-70b-versatile"),
        ("cerebras", "llama3.1-8b"),
        ("sambanova", "Meta-Llama-3.1-8B-Instruct"),
        ("pollinations", "openai-fast"),
    ]

    for prov_name, model_id in translate_chains:
        prov = ProviderFactory.get(prov_name, API_KEYS)
        if not prov or not await prov.health_check():
            continue
        try:
            resp = await prov.chat(messages, model_id, temperature=0.3, max_tokens=2048)
            if resp.success and resp.content.strip():
                log.info(f"🌐 Translated via {prov_name}/{model_id}")
                return resp.content.strip()
        except Exception as e:
            log.warning(f"Translate error {prov_name}: {e}")
            continue

    return f"[Translation failed] {text}"

# ============================================================
# URL FETCH — Universal URL Reader
# ============================================================

def _detect_platform(url: str) -> str:
    """Detect platform from URL"""
    url_lower = url.lower()
    if "twitter.com" in url_lower or "x.com" in url_lower:
        return "twitter"
    elif "instagram.com" in url_lower:
        return "instagram"
    elif "tiktok.com" in url_lower:
        return "tiktok"
    elif "youtube.com" in url_lower or "youtu.be" in url_lower:
        return "youtube"
    elif "reddit.com" in url_lower:
        return "reddit"
    elif "github.com" in url_lower:
        return "github"
    return "generic"

def _extract_youtube_id(url: str) -> Optional[str]:
    """Extract YouTube video ID from URL"""
    patterns = [
        r'(?:v=|/v/|youtu\.be/)([a-zA-Z0-9_-]{11})',
        r'(?:shorts/)([a-zA-Z0-9_-]{11})',
    ]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None

async def _fetch_via_jina(url: str) -> Optional[str]:
    """Fetch URL content via Jina Reader — works for most websites"""
    try:
        jina_url = f"https://r.jina.ai/{url}"
        headers = {
            "Accept": "text/markdown",
            "User-Agent": "Mozilla/5.0 (compatible; DiscordBot/1.0)"
        }
        async with aiohttp.ClientSession() as session:
            async with session.get(
                jina_url,
                headers=headers,
                timeout=aiohttp.ClientTimeout(total=20)
            ) as resp:
                if resp.status == 200:
                    content = await resp.text()
                    if len(content) > 8000:
                        content = content[:8000] + "\n\n[... content trimmed ...]"
                    log.info(f"📄 Jina Reader OK: {url[:60]}")
                    return content
                else:
                    log.warning(f"📄 Jina Reader HTTP {resp.status} for {url[:60]}")
    except Exception as e:
        log.warning(f"📄 Jina Reader error: {e}")
    return None

async def _fetch_via_bs4(url: str) -> Optional[str]:
    """Fallback: fetch with requests + BeautifulSoup"""
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        log.warning("beautifulsoup4 not installed")
        return None

    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
        }
        async with aiohttp.ClientSession() as session:
            async with session.get(
                url,
                headers=headers,
                timeout=aiohttp.ClientTimeout(total=15),
                allow_redirects=True
            ) as resp:
                if resp.status == 200:
                    html = await resp.text()
                    soup = BeautifulSoup(html, "html.parser")

                    for tag in soup(["script", "style", "nav", "footer", "header", "aside", "iframe"]):
                        tag.decompose()

                    article = soup.find("article")
                    if article:
                        text = article.get_text(separator="\n", strip=True)
                    else:
                        main = soup.find("main") or soup.find("body")
                        text = main.get_text(separator="\n", strip=True) if main else ""

                    lines = [line.strip() for line in text.split("\n") if line.strip()]
                    text = "\n".join(lines)

                    if len(text) > 6000:
                        text = text[:6000] + "\n\n[... content trimmed ...]"

                    if len(text) > 100:
                        log.info(f"📄 BS4 fetch OK: {url[:60]}")
                        return text
    except Exception as e:
        log.warning(f"📄 BS4 error: {e}")
    return None

async def _fetch_twitter(url: str) -> Optional[str]:
    """Fetch Twitter/X content via Nitter proxies or Jina"""
    nitter_instances = [
        "nitter.net",
        "nitter.privacydev.net",
        "nitter.poast.org",
    ]

    match = re.search(r'(?:twitter\.com|x\.com)/(\w+)/status/(\d+)', url)
    if not match:
        return await _fetch_via_jina(url)

    username, tweet_id = match.group(1), match.group(2)

    for instance in nitter_instances:
        try:
            nitter_url = f"https://{instance}/{username}/status/{tweet_id}"
            content = await _fetch_via_jina(nitter_url)
            if content and len(content) > 50:
                return f"[Tweet from @{username}]\n\n{content}"
        except:
            continue

    content = await _fetch_via_jina(url)
    if content:
        return f"[Tweet from @{username}]\n\n{content}"

    return None

async def _fetch_youtube_transcript(url: str) -> Optional[str]:
    """Fetch YouTube video info + transcript"""
    video_id = _extract_youtube_id(url)
    if not video_id:
        return None

    parts = []

    try:
        oembed_url = f"https://www.youtube.com/oembed?url=https://www.youtube.com/watch?v={video_id}&format=json"
        async with aiohttp.ClientSession() as session:
            async with session.get(oembed_url, timeout=aiohttp.ClientTimeout(total=10)) as resp:
                if resp.status == 200:
                    data = await resp.json()
                    parts.append(f"Title: {data.get('title', 'Unknown')}")
                    parts.append(f"Channel: {data.get('author_name', 'Unknown')}")
    except:
        pass

    try:
        import yt_dlp

        def _get_info():
            ydl_opts = {
                "quiet": True,
                "no_warnings": True,
                "skip_download": True,
                "writeautomaticsub": True,
                "subtitleslangs": ["id", "en"],
                "subtitlesformat": "json3",
            }
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                return ydl.extract_info(url, download=False)

        info = await asyncio.get_event_loop().run_in_executor(None, _get_info)

        if info:
            if not parts:
                parts.append(f"Title: {info.get('title', 'Unknown')}")
                parts.append(f"Channel: {info.get('channel', info.get('uploader', 'Unknown'))}")

            duration = info.get("duration", 0)
            if duration:
                mins, secs = divmod(duration, 60)
                parts.append(f"Duration: {int(mins)}:{int(secs):02d}")

            view_count = info.get("view_count")
            if view_count:
                parts.append(f"Views: {view_count:,}")

            desc = info.get("description", "")
            if desc:
                parts.append(f"Description: {desc[:500]}")

            subtitles = info.get("subtitles", {})
            auto_subs = info.get("automatic_captions", {})

            transcript_text = None

            for sub_source in [subtitles, auto_subs]:
                for lang in ["id", "en"]:
                    if lang in sub_source:
                        for fmt in sub_source[lang]:
                            if fmt.get("ext") == "json3":
                                try:
                                    sub_url = fmt["url"]
                                    async with aiohttp.ClientSession() as session:
                                        async with session.get(sub_url, timeout=aiohttp.ClientTimeout(total=10)) as resp:
                                            if resp.status == 200:
                                                sub_data = await resp.json()
                                                events = sub_data.get("events", [])
                                                words = []
                                                for event in events:
                                                    segs = event.get("segs", [])
                                                    for seg in segs:
                                                        w = seg.get("utf8", "").strip()
                                                        if w and w != "\n":
                                                            words.append(w)
                                                transcript_text = " ".join(words)
                                except:
                                    pass
                        if transcript_text:
                            break
                if transcript_text:
                    break

            if transcript_text:
                if len(transcript_text) > 5000:
                    transcript_text = transcript_text[:5000] + "... [trimmed]"
                parts.append(f"\nTranscript:\n{transcript_text}")

            log.info(f"🎬 YouTube info OK: {video_id}")
            return "\n".join(parts)

    except ImportError:
        log.warning("yt-dlp not installed, falling back to Jina")
    except Exception as e:
        log.warning(f"yt-dlp error: {e}")

    jina_content = await _fetch_via_jina(url)
    if jina_content:
        if parts:
            return "\n".join(parts) + "\n\n" + jina_content
        return jina_content

    return "\n".join(parts) if parts else None

async def _get_video_download_url(url: str) -> Optional[dict]:
    """Download video langsung pakai yt-dlp dengan Instagram proxy fallback"""

    original_url = url
    platform = _detect_platform(url)

    # ── INSTAGRAM: Coba proxy dulu karena sering butuh login ──
    if platform == "instagram":
        proxy_services = [
            ("ddinstagram.com", url.replace("instagram.com", "ddinstagram.com")),
            ("imginn.com", None),
        ]

        for service_name, proxy_url in proxy_services:
            if not proxy_url:
                continue
            log.info(f"📱 Trying Instagram proxy: {service_name}")
            result = await _try_ytdlp_download(proxy_url, original_url)
            if result:
                return result

        log.info("📱 Trying direct Instagram (may need cookies)")
        result = await _try_ytdlp_download(url, original_url)
        if result:
            return result

        return None

    # ── Platform lain: langsung yt-dlp ──
    return await _try_ytdlp_download(url, original_url)


async def _try_ytdlp_download(url: str, original_url: str = None) -> Optional[dict]:
    """Actual yt-dlp download attempt"""
    try:
        import yt_dlp

        def _download_direct():
            temp_dir = tempfile.mkdtemp()
            output_path = os.path.join(temp_dir, "video.mp4")

            ydl_opts = {
                "quiet": True,
                "no_warnings": True,
                "format": "bestvideo[height<=1440][ext=mp4]+bestaudio[ext=m4a]/best[height<=1440][ext=mp4]/best[ext=mp4]/best",
                "outtmpl": output_path,
                "merge_output_format": "mp4",
                "socket_timeout": 30,
                "retries": 3,
                "cookiefile": "/root/claw.ai/cookies.txt",
                "http_headers": {
                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
                    "Accept-Language": "en-US,en;q=0.5",
                },
            }

            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                info = ydl.extract_info(url, download=True)

                actual_path = output_path
                if not os.path.exists(actual_path):
                    for f in os.listdir(temp_dir):
                        actual_path = os.path.join(temp_dir, f)
                        break

                if not os.path.exists(actual_path):
                    return None

                title = info.get("title", "video")[:50]
                clean_title = re.sub(r'[^\w\s-]', '', title).strip()
                clean_title = re.sub(r'\s+', '_', clean_title) or "video"

                return {
                    "local_path": actual_path,
                    "temp_dir": temp_dir,
                    "filename": f"{clean_title}.mp4",
                    "status": "local",
                    "filesize": os.path.getsize(actual_path),
                    "title": info.get("title", ""),
                    "uploader": info.get("uploader", info.get("channel", "")),
                    "duration": info.get("duration", 0),
                    "original_url": original_url or url,
                }

        result = await asyncio.get_event_loop().run_in_executor(None, _download_direct)
        if result and result.get("local_path") and os.path.exists(result["local_path"]):
            log.info(f"🎬 yt-dlp OK: {result['filename']} ({result.get('filesize', 0) / 1_000_000:.1f}MB)")
            return result

    except Exception as e:
        log.warning(f"yt-dlp error: {e}")

    return None

async def do_fetch_url(url: str, action: str = "read") -> str:
    """Universal URL fetcher — read content or get download URL"""
    platform = _detect_platform(url)

    log.info(f"📄 Fetching URL: {url[:80]} | platform={platform} | action={action}")

    # ── DOWNLOAD ACTION ──
    if action == "download":
        download_info = await _get_video_download_url(url)
        if download_info and download_info.get("status") == "local":
            return json.dumps({
                "type": "download",
                "local_path": download_info["local_path"],
                "temp_dir": download_info["temp_dir"],
                "filename": download_info.get("filename", "video.mp4"),
                "platform": platform,
                "original_url": url,
                "method": "local",
                "title": download_info.get("title", ""),
                "uploader": download_info.get("uploader", ""),
            })
        return "Cannot download video from this URL. The content might be protected, require login, or not a video."

    # ── READ/SUMMARIZE ACTION ──
    content = None

    if platform == "twitter":
        content = await _fetch_twitter(url)

    elif platform == "youtube":
        content = await _fetch_youtube_transcript(url)

    elif platform in ("tiktok", "instagram", "reddit"):
        content = await _fetch_via_jina(url)

    elif platform == "github":
        content = await _fetch_via_jina(url)

    # Generic / fallback
    if not content:
        content = await _fetch_via_jina(url)

    if not content:
        content = await _fetch_via_bs4(url)

    if not content:
        return (
            f"Cannot access content from {url}. "
            f"Possible reasons: website requires login, anti-bot protection, "
            f"or content is not publicly accessible."
        )

    return f"[Content from {platform}: {url}]\n\n{content}"


# ============================================================
# TOOL DEFINITIONS (8 Tools)
# ============================================================

WEB_SEARCH_TOOL = {
    "type": "function",
    "function": {
        "name": "web_search",
        "description": (
            "Search the internet for real-time information. "
            "Use this when asked about: current events, news, prices, "
            "who is president/leader/CEO, sports results, "
            "anything that might have changed recently, or any fact "
            "you are not confident about."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "Search query in the most relevant language"}
            },
            "required": ["query"]
        }
    }
}

GET_TIME_TOOL = {
    "type": "function",
    "function": {
        "name": "get_time",
        "description": (
            "Get current date and time in a specific timezone. "
            "Use for questions about: what time is it, today's date, "
            "what day is it, current time in a city/country."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "timezone": {
                    "type": "string",
                    "description": "Timezone like Asia/Jakarta, America/New_York, Europe/London, Asia/Tokyo. Default: Asia/Jakarta"
                }
            },
            "required": []
        }
    }
}

GET_WEATHER_TOOL = {
    "type": "function",
    "function": {
        "name": "get_weather",
        "description": (
            "Get current weather for a city. "
            "Use when asked about weather, temperature, rain, humidity in a location."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "city": {
                    "type": "string",
                    "description": "City name, e.g. Jakarta, Tokyo, London, New York"
                }
            },
            "required": ["city"]
        }
    }
}

GET_FORECAST_TOOL = {
    "type": "function",
    "function": {
        "name": "get_forecast",
        "description": (
            "Get weather forecast for the next few days. "
            "Use when asked about: tomorrow's weather, weekly forecast, "
            "will it rain tomorrow, cuaca besok, ramalan cuaca."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "city": {
                    "type": "string",
                    "description": "City name, e.g. Jakarta, Tokyo, London"
                },
                "days": {
                    "type": "integer",
                    "description": "Number of forecast days (1-7). Default: 3"
                }
            },
            "required": ["city"]
        }
    }
}

CALCULATE_TOOL = {
    "type": "function",
    "function": {
        "name": "calculate",
        "description": (
            "Perform mathematical calculations. "
            "Use for math questions, unit conversions, percentages, etc."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "expression": {
                    "type": "string",
                    "description": "Math expression like '5 + 3 * 2', 'sqrt(144)', '15% of 200', '2**10'"
                }
            },
            "required": ["expression"]
        }
    }
}

TRANSLATE_TOOL = {
    "type": "function",
    "function": {
        "name": "translate",
        "description": (
            "Translate text between languages naturally like a native speaker. "
            "Understands slang, idioms, and cultural context. "
            "Use when user asks to translate something, or wants to know how to say something in another language."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "text": {"type": "string", "description": "The text to translate"},
                "target_language": {"type": "string", "description": "Target language, e.g. English, Indonesian, Japanese, Korean, Spanish"},
                "style": {
                    "type": "string",
                    "enum": ["natural", "formal", "casual"],
                    "description": "Translation style. natural=everyday speech, formal=polished, casual=slang/texting. Default: natural"
                }
            },
            "required": ["text", "target_language"]
        }
    }
}

PLAY_MUSIC_TOOL = {
    "type": "function",
    "function": {
        "name": "play_music",
        "description": (
            "Control music playback in Discord voice channel. "
            "Use when user wants to: play a song, skip track, stop/disconnect music, "
            "pause, or resume playback. "
            "If user is NOT in a voice channel, tell them to join one first."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "action": {
                    "type": "string",
                    "enum": ["play", "skip", "stop", "pause", "resume"],
                    "description": "Music action"
                },
                "query": {
                    "type": "string",
                    "description": "Song name, artist, or URL. Required for 'play' action."
                }
            },
            "required": ["action"]
        }
    }
}

FETCH_URL_TOOL = {
    "type": "function",
    "function": {
        "name": "fetch_url",
        "description": (
            "Read, summarize, or download content from ANY URL. "
            "\n\n"
            "SUPPORTED PLATFORMS FOR DOWNLOAD:\n"
            "• TikTok (no watermark) ✅\n"
            "• YouTube (videos, shorts, music) ✅\n"
            "• Twitter/X (videos, GIFs) ✅\n"
            "• Reddit (videos) ✅\n"
            "• Facebook (videos, reels) ✅\n"
            "• Instagram (reels, posts) ⚠️ May require fallback\n"
            "• And 1000+ other sites\n"
            "\n"
            "SUPPORTED FOR READING/SUMMARIZING:\n"
            "• News articles, blogs, Medium\n"
            "• Twitter/X tweets, YouTube info\n"
            "• GitHub README, any website\n"
            "\n"
            "ACTIONS:\n"
            "• action='download' → Download video/audio\n"
            "• action='read' → Read content\n"
            "• action='summarize' → Read for summary\n"
            "\n"
            "NOTE: Instagram sometimes requires login. If download fails, "
            "inform the user and suggest they try a different link or use a web-based downloader."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "url": {"type": "string", "description": "The full URL"},
                "action": {
                    "type": "string",
                    "enum": ["read", "summarize", "download"],
                    "description": "Action to perform"
                }
            },
            "required": ["url"]
        }
    }
}

GENERATE_IMAGE_TOOL = {
    "type": "function",
    "function": {
        "name": "generate_image",
        "description": (
            "Generate an image from text description using AI. "
            "Use when user asks to: create/generate/make/draw an image, picture, photo, artwork. "
            "Describe the image in English for best results."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "prompt": {
                    "type": "string",
                    "description": "Image description in English. Be detailed for better results. Example: 'a cute cat wearing astronaut suit floating in space, digital art, high quality'"
                },
                "size": {
                    "type": "string",
                    "enum": ["square", "wide", "tall"],
                    "description": "Image aspect ratio. square=1:1, wide=16:9, tall=9:16. Default: square"
                }
            },
            "required": ["prompt"]
        }
    }
}

CREATE_DOCUMENT_TOOL = {
    "type": "function",
    "function": {
        "name": "create_document",
        "description": (
            "Create office documents or code files and send to user. "
            "\n\n"
            "SUPPORTED FILE TYPES:\n"
            "• Word (.docx) — letters, reports, formal documents\n"
            "• Excel (.xlsx) — tables, data, spreadsheets, schedules\n"
            "• Code files (.py, .js, .html, .css, .json, .sql, etc)\n"
            "• Text files (.txt, .md, .csv)\n"
            "\n"
            "IMPORTANT FOR EXCEL:\n"
            "You MUST provide content as JSON array of arrays. "
            "First row is header. Example:\n"
            '[[\"Name\",\"Age\",\"City\"],[\"John\",25,\"NYC\"],[\"Jane\",30,\"LA\"]]'
            "\n\n"
            "IMPORTANT FOR WORD:\n"
            "Use markdown-like formatting:\n"
            "# Heading 1, ## Heading 2, ### Heading 3\n"
            "- bullet point, 1. numbered list\n"
            "Regular text as paragraphs."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "file_type": {
                    "type": "string",
                    "enum": ["docx", "xlsx", "py", "js", "html", "css", "json", "txt", "md", "sql", "csv", "java", "cpp", "sh"],
                    "description": "File type to create"
                },
                "filename": {
                    "type": "string",
                    "description": "Output filename with extension, e.g. 'report.docx', 'data.xlsx', 'script.py'"
                },
                "title": {
                    "type": "string",
                    "description": "Document title (for Word/Excel). Optional."
                },
                "content": {
                    "type": "string",
                    "description": "File content. For Excel: JSON array of arrays. For Word: markdown-formatted text. For code: raw code."
                }
            },
            "required": ["file_type", "filename", "content"]
        }
    }
}


GET_SERVER_INFO_TOOL = {
    "type": "function",
    "function": {
        "name": "get_server_info",
        "description": (
            "Get Discord server information: members, channels, voice channel users. "
            "Use when user asks: who is online, siapa yang ada di server, "
            "siapa di voice channel, list members, server info."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "info_type": {
                    "type": "string",
                    "enum": ["members", "voice", "channels", "all"],
                    "description": "Type of info: members=member list, voice=who in voice channels, channels=channel list, all=everything"
                }
            },
            "required": []
        }
    }
}

SYSTEM_STATUS_TOOL = {
    "type": "function",
    "function": {
        "name": "system_status",
        "description": (
            "Get bot's own system status: CPU, memory, disk usage, service status, recent errors. "
            "Use when user asks about: bot status, is bot okay, any errors, system health, "
            "why is bot slow, check logs, bot uptime, dll."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "check_type": {
                    "type": "string",
                    "enum": ["full", "errors", "logs", "resources", "git"],
                    "description": "Type of check. full=complete report, errors=recent errors only, logs=recent logs, resources=CPU/memory/disk, git=git status"
                },
                "log_lines": {
                    "type": "integer",
                    "description": "Number of log lines to fetch (default: 30)"
                }
            },
            "required": []
        }
    }
}

READ_SOURCE_TOOL = {
    "type": "function",
    "function": {
        "name": "read_source",
        "description": (
            "Read bot's own source code file for debugging. "
            "Use when user asks to check code, debug issue, or see how something works."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "File to read: main.py, config.py, core/handler.py, core/providers.py, etc."
                }
            },
            "required": ["filename"]
        }
    }
}

BOT_CONTROL_TOOL = {
    "type": "function",
    "function": {
        "name": "bot_control",
        "description": (
            "Control bot service: restart, update from git. "
            "DANGEROUS - only use when explicitly requested by admin. "
            "Use when user says: restart bot, update bot, pull latest code."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "action": {
                    "type": "string",
                    "enum": ["restart", "git_pull", "git_status"],
                    "description": "Action to perform"
                }
            },
            "required": ["action"]
        }
    }
}

SET_REMINDER_TOOL = {
    "type": "function",
    "function": {
        "name": "set_reminder",
        "description": (
            "Set a reminder or scheduled notification. "
            "Use when user wants to be reminded about something at a specific time or after X minutes. "
            "\n"
            "Examples:\n"
            "• 'Ingatkan aku meeting jam 3' → daily_time='15:00', message='Meeting'\n"
            "• 'Reminder 30 menit lagi' → minutes=30, message='Reminder'\n"
            "• 'Set alarm sahur jam 3 pagi' → daily_time='03:00', message='Sahur'\n"
            "• 'Ingatkan imsak jam 4:15' → daily_time='04:15', message='Imsak'\n"
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "message": {
                    "type": "string",
                    "description": "Reminder message"
                },
                "minutes": {
                    "type": "integer",
                    "description": "Remind after X minutes from now. Use this OR daily_time, not both."
                },
                "daily_time": {
                    "type": "string",
                    "description": "Time in HH:MM format (24h). For recurring daily reminder. e.g. '03:00' for sahur, '04:15' for imsak"
                },
                "recurring": {
                    "type": "boolean",
                    "description": "If true, reminder repeats daily at the same time. Default: false"
                },
                "timezone": {
                    "type": "string",
                    "description": "Timezone. Default: Asia/Jakarta"
                }
            },
            "required": ["message"]
        }
    }
}

SEND_MESSAGE_TOOL = {
    "type": "function",
    "function": {
        "name": "send_message",
        "description": (
            "Send a message or file to a different channel or DM the user. "
            "Use when user asks to: send to another channel, DM me, kirim ke DM, "
            "send to #channel-name, etc."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "destination": {
                    "type": "string",
                    "enum": ["dm", "channel"],
                    "description": "Where to send: 'dm' for Direct Message, 'channel' for another channel"
                },
                "channel_name": {
                    "type": "string",
                    "description": "Channel name (without #) if destination is 'channel'. e.g. 'general', 'bot-spam'"
                },
                "message": {
                    "type": "string",
                    "description": "Message content to send"
                }
            },
            "required": ["destination"]
        }
    }
}

TOOLS_LIST = [
    WEB_SEARCH_TOOL, GET_TIME_TOOL, GET_WEATHER_TOOL, GET_FORECAST_TOOL,
    CALCULATE_TOOL, TRANSLATE_TOOL, PLAY_MUSIC_TOOL, FETCH_URL_TOOL,
    GENERATE_IMAGE_TOOL, CREATE_DOCUMENT_TOOL, SET_REMINDER_TOOL,
    SEND_MESSAGE_TOOL, GET_SERVER_INFO_TOOL,  # <-- TAMBAH
    SYSTEM_STATUS_TOOL, READ_SOURCE_TOOL, BOT_CONTROL_TOOL
]


# ============================================================
# MODE DETECTOR
# ============================================================

class ModeDetector:
    SEARCH_KW = ["berita terbaru", "harga sekarang", "news today", "current price", "latest news"]
    REASON_KW = ["jelaskan step by step", "hitung ", "analisis ", "solve ", "tulis kode", "write code"]

    @classmethod
    def detect(cls, content):
        lower = content.lower()
        for kw in cls.SEARCH_KW:
            if kw in lower: return "search"
        for kw in cls.REASON_KW:
            if kw in lower: return "reasoning"
        return "normal"


# ============================================================
# TOOL CALL EXECUTOR — 8 Tools
# ============================================================

async def execute_tool_call(tool_name: str, tool_args: dict) -> str:
    """Eksekusi tool yang diminta AI"""

    # ── WEB SEARCH ──
    if tool_name == "web_search":
        query = tool_args.get("query", "")
        log.info(f"🔍 Tool: web_search({query})")
        return await do_search(query)

    # ── GET TIME ──
    elif tool_name == "get_time":
        from skills.time_skill import get_current_time
        tz = tool_args.get("timezone", "Asia/Jakarta")
        log.info(f"🕐 Tool: get_time({tz})")
        result = get_current_time(tz)
        if result["success"]:
            return f"Current time in {tz}: {result['full']}"
        return f"Error: {result.get('error', 'Unknown timezone')}"

    # ── GET WEATHER ──
    elif tool_name == "get_weather":
        from skills.weather_skill import get_weather
        city = tool_args.get("city", "Jakarta")
        log.info(f"🌤️ Tool: get_weather({city})")
        result = await get_weather(city)
        if result["success"]:
            response = (
                f"Cuaca di {result['city']}: {result['description']}\n"
                f"🌡️ Suhu: {result['temp']}°C (terasa {result['feels_like']}°C)\n"
                f"💧 Kelembapan: {result['humidity']}%\n"
                f"💨 Angin: {result['wind_speed']} km/h"
            )
            if result.get('pressure'):
                response += f"\n🌡️ Tekanan: {result['pressure']} hPa"
            if result.get('visibility'):
                response += f"\n👁️ Visibilitas: {result['visibility']} km"
            response += f"\n\n📡 Sumber: {result.get('source', 'Weather API')}"
            return response
        return f"Error: {result.get('error', 'Kota tidak ditemukan')}"

    # ── GET FORECAST ──
    elif tool_name == "get_forecast":
        from skills.weather_skill import get_forecast
        city = tool_args.get("city", "Jakarta")
        days = tool_args.get("days", 3)
        log.info(f"🌤️ Tool: get_forecast({city}, {days} days)")
        result = await get_forecast(city, days)
        if result["success"]:
            lines = [f"📅 Ramalan Cuaca {result['city']} ({days} hari):\n"]
            for f in result["forecasts"]:
                lines.append(
                    f"• {f['date']}: {f['description']}\n"
                    f"  🌡️ {f['temp_min']}°C - {f['temp_max']}°C | "
                    f"🌧️ {f['rain_chance']}% hujan | "
                    f"💨 {f['wind_max']} km/h"
                )
            return "\n".join(lines)
        return f"Error: {result.get('error', 'Kota tidak ditemukan')}"

    # ── CALCULATE ──
    elif tool_name == "calculate":
        expression = tool_args.get("expression", "")
        log.info(f"🔢 Tool: calculate({expression})")
        try:
            expr = expression.replace("^", "**").replace("×", "*").replace("÷", "/")
            pct_match = re.match(r'([\d.]+)%\s*of\s*([\d.]+)', expr, re.IGNORECASE)
            if pct_match:
                pct, val = float(pct_match.group(1)), float(pct_match.group(2))
                result = pct / 100 * val
                return f"{expression} = {result}"
            allowed = {
                "sqrt": math.sqrt, "sin": math.sin, "cos": math.cos,
                "tan": math.tan, "log": math.log, "log10": math.log10,
                "log2": math.log2, "ceil": math.ceil, "floor": math.floor,
                "abs": abs, "round": round, "pow": pow, "max": max, "min": min,
                "pi": math.pi, "e": math.e,
            }
            result = eval(expr, {"__builtins__": {}}, allowed)
            return f"{expression} = {result}"
        except Exception as e:
            return f"Calculation error: {e}"

    # ── TRANSLATE ──
    elif tool_name == "translate":
        text = tool_args.get("text", "")
        target = tool_args.get("target_language", "English")
        style = tool_args.get("style", "natural")
        log.info(f"🌐 Tool: translate(→{target}, style={style})")
        return await do_translate(text, target, style)

    # ── PLAY MUSIC ──
    elif tool_name == "play_music":
        action = tool_args.get("action", "play")
        query = tool_args.get("query", "")
        log.info(f"🎵 Tool: play_music(action={action}, query={query})")
        if action == "play":
            return f"Music command accepted. Now playing '{query}' in the user's voice channel."
        elif action == "skip":
            return "Skipping to the next track."
        elif action == "stop":
            return "Stopping music and disconnecting from voice channel."
        elif action == "pause":
            return "Pausing current track."
        elif action == "resume":
            return "Resuming playback."
        return f"Music action '{action}' accepted."

    # ── FETCH URL ──
    elif tool_name == "fetch_url":
        url = tool_args.get("url", "")
        action = tool_args.get("action", "read")
        log.info(f"📄 Tool: fetch_url({url[:60]}, action={action})")
        return await do_fetch_url(url, action)

    # ── GENERATE IMAGE ──
    elif tool_name == "generate_image":
        prompt = tool_args.get("prompt", "")
        size = tool_args.get("size", "square")
        log.info(f"🖼️ Tool: generate_image(prompt={prompt[:50]}, size={size})")

        size_map = {
            "square": (1024, 1024),
            "wide": (1280, 720),
            "tall": (720, 1280),
        }
        w, h = size_map.get(size, (1024, 1024))

        image_url = f"https://image.pollinations.ai/prompt/{prompt}?width={w}&height={h}&nologo=true&seed={int(datetime.now().timestamp())}"

        return json.dumps({
            "type": "image",
            "image_url": image_url,
            "prompt": prompt,
            "size": size,
        })
        
            # ── CREATE DOCUMENT ──
    elif tool_name == "create_document":
        file_type = tool_args.get("file_type", "txt")
        filename = tool_args.get("filename", f"document.{file_type}")
        title = tool_args.get("title", "")
        content = tool_args.get("content", "")
        log.info(f"📄 Tool: create_document({filename})")
        
        result = await create_document(file_type, filename, content, title)
        if result:
            return json.dumps(result)
        return f"Failed to create {filename}"

        # ── SET REMINDER ──
    elif tool_name == "set_reminder":
        message = tool_args.get("message", "Reminder!")
        minutes = tool_args.get("minutes")
        daily_time = tool_args.get("daily_time")
        recurring = tool_args.get("recurring", False)
        timezone = tool_args.get("timezone", "Asia/Jakarta")
        log.info(f"⏰ Tool: set_reminder(msg={message}, min={minutes}, time={daily_time})")
        
        import pytz
        tz = pytz.timezone(timezone)
        now = datetime.now(tz)
        
        if minutes:
            trigger = now + timedelta(minutes=minutes)
            return json.dumps({
                "type": "reminder",
                "message": message,
                "trigger_minutes": minutes,
                "trigger_time": trigger.strftime("%H:%M:%S"),
                "recurring": False,
                "timezone": timezone,
            })
        elif daily_time:
            return json.dumps({
                "type": "reminder",
                "message": message,
                "daily_time": daily_time,
                "recurring": recurring,
                "timezone": timezone,
                "trigger_time": daily_time,
            })
        else:
            return json.dumps({
                "type": "reminder",
                "message": message,
                "trigger_minutes": 5,
                "trigger_time": (now + timedelta(minutes=5)).strftime("%H:%M:%S"),
                "recurring": False,
                "timezone": timezone,
            })

    # ══════════════════════════════════════════
    # TARUH DI SINI ↓↓↓
    # ══════════════════════════════════════════

    # ── SEND MESSAGE ──
    elif tool_name == "send_message":
        destination = tool_args.get("destination", "dm")
        channel_name = tool_args.get("channel_name", "")
        msg_content = tool_args.get("message", "")
        
        log.info(f"📤 Tool: send_message(dest={destination}, channel={channel_name})")
        
        return json.dumps({
            "type": "send_message",
            "destination": destination,
            "channel_name": channel_name,
            "message": msg_content,
        })

    # ══════════════════════════════════════════
    # AKHIR TAMBAHAN ↑↑↑
    # ══════════════════════════════════════════


    # ── GET SERVER INFO ──
    elif tool_name == "get_server_info":
        info_type = tool_args.get("info_type", "all")
        server_info = tool_args.get("_server_info", {})
        log.info(f"👥 Tool: get_server_info({info_type}) | has_data={bool(server_info)} | members={len(server_info.get('all_members', []))}")
        
        if not server_info:
            return "Server info tidak tersedia. Bot mungkin tidak memiliki akses ke member list."
        
        lines = []
        server_name = server_info.get("server_name", "Unknown")
        total = server_info.get("total_members", 0)
        lines.append(f"Server: {server_name} | Total: {total} members")
        
        if info_type in ("members", "all"):
            all_members = server_info.get("all_members", [])
            online = server_info.get("online_members", [])
            
            if all_members:
                lines.append(f"\nSemua Member ({len(all_members)} orang, non-bot):")
                for m in all_members:
                    lines.append(f"  {m}")
            
            if online:
                lines.append(f"\nYang Online ({len(online)}):")
                for m in online:
                    lines.append(f"  {m}")
            else:
                lines.append("\nTidak ada member yang online saat ini.")
        
        if info_type in ("voice", "all"):
            voice = server_info.get("voice_channels", [])
            lines.append("\nVoice Channels:")
            if voice:
                for v in voice:
                    lines.append(f"  {v}")
            else:
                lines.append("  Tidak ada voice channel.")
        
        if info_type in ("channels", "all"):
            channels = server_info.get("text_channels", [])
            lines.append(f"\nText Channels ({len(channels)}):")
            for ch in channels:
                lines.append(f"  {ch}")
        
        result = "\n".join(lines)
        log.info(f"👥 Returning: {result[:200]}")
        return result

    # ── SYSTEM STATUS ──
    elif tool_name == "system_status":
        from skills.system_skill import (
            format_system_report, format_error_report, 
            get_recent_logs, get_system_status, get_git_status
        )
        
        check_type = tool_args.get("check_type", "full")
        log_lines = tool_args.get("log_lines", 30)
        
        log.info(f"🖥️ Tool: system_status({check_type})")
        
        if check_type == "full":
            return format_system_report()
        elif check_type == "errors":
            return format_error_report()
        elif check_type == "logs":
            result = get_recent_logs(lines=log_lines)
            if result["success"]:
                return f"Recent logs ({result['total_lines']} lines):\n```\n{result['log_text']}\n```"
            return f"Failed: {result['error']}"
        elif check_type == "resources":
            result = get_system_status()
            if result["success"]:
                return (
                    f"CPU: {result['cpu_percent']}%\n"
                    f"Memory: {result['memory_used_gb']}/{result['memory_total_gb']} GB ({result['memory_percent']}%)\n"
                    f"Disk: {result['disk_used_gb']}/{result['disk_total_gb']} GB ({result['disk_percent']}%)\n"
                    f"Uptime: {result['uptime']}"
                )
            return f"Failed: {result['error']}"
        elif check_type == "git":
            result = get_git_status()
            if result["success"]:
                return (
                    f"Branch: {result['branch']}\n"
                    f"Last commit: {result['last_commit']}\n"
                    f"Has updates: {'Yes' if result['has_updates'] else 'No'}\n"
                    f"Status: {result['status']}"
                )
            return f"Failed: {result['error']}"
        
        return format_system_report()

    # ── READ SOURCE ──
    elif tool_name == "read_source":
        from skills.system_skill import read_bot_file
        
        filename = tool_args.get("filename", "main.py")
        log.info(f"📄 Tool: read_source({filename})")
        
        result = read_bot_file(filename)
        if result["success"]:
            return f"File: {result['filename']} ({result['size_bytes']} bytes)\n```python\n{result['content']}\n```"
        return f"Error: {result['error']}"

    # ── BOT CONTROL ──
    elif tool_name == "bot_control":
        from skills.system_skill import restart_bot_service, git_pull_update, get_git_status
        
        action = tool_args.get("action", "git_status")
        log.info(f"🔧 Tool: bot_control({action})")
        
        if action == "git_status":
            result = get_git_status()
            if result["success"]:
                return (
                    f"Branch: {result['branch']}\n"
                    f"Last commit: {result['last_commit']}\n"
                    f"Updates available: {'Yes ⚠️' if result['has_updates'] else 'No ✅'}"
                )
            return f"Error: {result['error']}"
        
        elif action == "git_pull":
            result = git_pull_update()
            if result["success"]:
                return f"✅ Git pull successful!\n{result['output']}\n\n⚠️ Restart bot to apply changes."
            return f"❌ Git pull failed: {result['error']}"
        
        elif action == "restart":
            result = restart_bot_service()
            if result["success"]:
                return "🔄 Bot is restarting... (this message may not be sent)"
            return f"❌ Restart failed: {result['error']}"
        
        return f"Unknown action: {action}"

    # ============ AKHIR TAMBAHAN ============
        
    return f"Unknown tool: {tool_name}"


# ============================================================
# FILE READER — Read uploaded files
# ============================================================

async def read_uploaded_file(file_url: str, filename: str) -> Optional[str]:
    """Download and read content from Discord attachment"""
    try:
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        
        # ── IMAGE FILES — Return URL for vision processing ──
        image_extensions = {"jpg", "jpeg", "png", "gif", "webp", "bmp"}
        
        if ext in image_extensions:
            log.info(f"🖼️ Image detected: {filename}")
            return json.dumps({
                "type": "image_attachment",
                "url": file_url,
                "filename": filename,
            })
        # ============ AKHIR TAMBAHAN ============
        
        async with aiohttp.ClientSession() as session:
            async with session.get(file_url, timeout=aiohttp.ClientTimeout(total=30)) as resp:
                if resp.status != 200:
                    return None
                data = await resp.read()
        
        # ── TEXT / CODE FILES ──
        text_extensions = {
            "py", "js", "ts", "html", "css", "json", "xml", "yaml", "yml",
            "txt", "md", "csv", "log", "env", "ini", "cfg", "conf",
            "sh", "bash", "bat", "ps1", "sql", "r", "rb", "php",
            "java", "kt", "swift", "go", "rs", "c", "cpp", "h", "hpp",
            "vue", "svelte", "jsx", "tsx", "dart", "lua", "toml",
        }
        
        if ext in text_extensions:
            try:
                text = data.decode("utf-8")
            except UnicodeDecodeError:
                text = data.decode("latin-1")
            
            if len(text) > 15000:
                text = text[:15000] + "\n\n[... file trimmed, too long ...]"
            
            log.info(f"📄 Read text file: {filename} ({len(text)} chars)")
            return f"[File: {filename}]\n```{ext}\n{text}\n```"
        
        # ── PDF ──
        elif ext == "pdf":
            try:
                from pypdf import PdfReader
                import io
                reader = PdfReader(io.BytesIO(data))
                pages = []
                for i, page in enumerate(reader.pages[:20]):  # Max 20 pages
                    text = page.extract_text()
                    if text:
                        pages.append(f"--- Page {i+1} ---\n{text}")
                
                full_text = "\n\n".join(pages)
                if len(full_text) > 15000:
                    full_text = full_text[:15000] + "\n\n[... PDF trimmed ...]"
                
                log.info(f"📄 Read PDF: {filename} ({len(reader.pages)} pages)")
                return f"[PDF: {filename} — {len(reader.pages)} pages]\n{full_text}"
            except ImportError:
                return f"[PDF: {filename}] — pypdf not installed, cannot read PDF"
            except Exception as e:
                return f"[PDF: {filename}] — Error reading: {e}"
        
        # ── DOCX ──
        elif ext == "docx":
            try:
                from docx import Document
                import io
                doc = Document(io.BytesIO(data))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                full_text = "\n".join(paragraphs)
                
                if len(full_text) > 15000:
                    full_text = full_text[:15000] + "\n\n[... document trimmed ...]"
                
                log.info(f"📄 Read DOCX: {filename} ({len(paragraphs)} paragraphs)")
                return f"[Word Document: {filename}]\n{full_text}"
            except ImportError:
                return f"[DOCX: {filename}] — python-docx not installed"
            except Exception as e:
                return f"[DOCX: {filename}] — Error reading: {e}"
        
        # ── XLSX ──
        elif ext in ("xlsx", "xls"):
            try:
                from openpyxl import load_workbook
                import io
                wb = load_workbook(io.BytesIO(data), read_only=True)
                sheets_text = []
                
                for sheet_name in wb.sheetnames[:5]:  # Max 5 sheets
                    ws = wb[sheet_name]
                    rows = []
                    for row in ws.iter_rows(max_row=100, values_only=True):  # Max 100 rows
                        row_data = [str(cell) if cell is not None else "" for cell in row]
                        rows.append(" | ".join(row_data))
                    
                    if rows:
                        sheets_text.append(f"[Sheet: {sheet_name}]\n" + "\n".join(rows))
                
                full_text = "\n\n".join(sheets_text)
                if len(full_text) > 15000:
                    full_text = full_text[:15000] + "\n\n[... spreadsheet trimmed ...]"
                
                log.info(f"📄 Read XLSX: {filename} ({len(wb.sheetnames)} sheets)")
                return f"[Excel: {filename}]\n{full_text}"
            except ImportError:
                return f"[XLSX: {filename}] — openpyxl not installed"
            except Exception as e:
                return f"[XLSX: {filename}] — Error reading: {e}"
        
        else:
            return f"[File: {filename}] — Unsupported file type: .{ext}"
    
    except Exception as e:
        log.error(f"📄 File read error: {e}")
        return f"[File: {filename}] — Error: {e}"


# ============================================================
# FILE CREATOR — Generate office documents
# ============================================================

async def create_document(file_type: str, filename: str, content: str, title: str = "") -> Optional[dict]:
    """Create Word, Excel, or Code file"""
    temp_dir = tempfile.mkdtemp()
    
    try:
        # ── WORD (.docx) ──
        if file_type == "docx":
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            if title:
                heading = doc.add_heading(title, level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Parse content — support basic formatting
            lines = content.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    doc.add_paragraph("")
                elif line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("- ") or line.startswith("• "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                elif re.match(r'^\d+\.\s', line):
                    doc.add_paragraph(re.sub(r'^\d+\.\s', '', line), style="List Number")
                else:
                    doc.add_paragraph(line)
            
            filepath = os.path.join(temp_dir, filename)
            doc.save(filepath)
            log.info(f"📄 Created DOCX: {filename}")
            
            return {
                "type": "upload_file",
                "local_path": filepath,
                "temp_dir": temp_dir,
                "filename": filename,
            }
        
        # ── EXCEL (.xlsx) ──
        elif file_type == "xlsx":
            from openpyxl import Workbook
            from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
            
            wb = Workbook()
            ws = wb.active
            ws.title = title or "Sheet1"
            
            # Parse content as JSON array of arrays
            try:
                rows_data = json.loads(content)
                if not isinstance(rows_data, list):
                    rows_data = [[content]]
            except (json.JSONDecodeError, TypeError):
                # Fallback: parse as text table
                rows_data = []
                for line in content.strip().split("\n"):
                    if line.strip():
                        # Split by | or , or tab
                        if "|" in line:
                            cells = [c.strip() for c in line.split("|") if c.strip()]
                        elif "\t" in line:
                            cells = [c.strip() for c in line.split("\t")]
                        elif "," in line:
                            cells = [c.strip() for c in line.split(",")]
                        else:
                            cells = [line.strip()]
                        rows_data.append(cells)
            
            # Write data to worksheet
            for row_idx, row in enumerate(rows_data, 1):
                for col_idx, value in enumerate(row, 1):
                    cell = ws.cell(row=row_idx, column=col_idx, value=value)
                    
                    # Style header row (first row)
                    if row_idx == 1:
                        cell.font = Font(bold=True, color="FFFFFF", size=11)
                        cell.fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
                        cell.alignment = Alignment(horizontal="center", vertical="center")
                    else:
                        cell.alignment = Alignment(vertical="center")
                    
                    # Add border
                    thin_border = Border(
                        left=Side(style="thin"),
                        right=Side(style="thin"),
                        top=Side(style="thin"),
                        bottom=Side(style="thin"),
                    )
                    cell.border = thin_border
            
            # Auto-adjust column width
            for col in ws.columns:
                max_length = 0
                col_letter = col[0].column_letter
                for cell in col:
                    try:
                        if cell.value:
                            max_length = max(max_length, len(str(cell.value)))
                    except:
                        pass
                ws.column_dimensions[col_letter].width = min(max_length + 4, 50)
            
            filepath = os.path.join(temp_dir, filename)
            wb.save(filepath)
            log.info(f"📊 Created XLSX: {filename}")
            
            return {
                "type": "upload_file",
                "local_path": filepath,
                "temp_dir": temp_dir,
                "filename": filename,
            }
        
        # ── CODE / TEXT FILES ──
        elif file_type in ("py", "js", "ts", "html", "css", "json", "txt", "md", 
                           "sql", "sh", "yaml", "xml", "csv", "java", "cpp", "c",
                           "go", "rs", "rb", "php", "swift", "kt", "dart", "lua"):
            filepath = os.path.join(temp_dir, filename)
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(content)
            
            log.info(f"💻 Created code file: {filename}")
            
            return {
                "type": "upload_file",
                "local_path": filepath,
                "temp_dir": temp_dir,
                "filename": filename,
            }
        
        else:
            log.warning(f"Unsupported file type: {file_type}")
            shutil.rmtree(temp_dir, ignore_errors=True)
            return None
    
    except Exception as e:
        log.error(f"📄 Create document error: {e}")
        shutil.rmtree(temp_dir, ignore_errors=True)
        return None


# ============================================================
# TOOL CALLING HANDLER
# ============================================================

async def handle_with_tools(messages: list, prov_name: str, model: str,
                             guild_id: int = 0, settings: dict = None) -> tuple:
    """
    Returns: (AIResponse, note_string, actions_list)
    """
    from core.providers import supports_tool_calling

    if not supports_tool_calling(prov_name):
        return None, None, []

    prov = ProviderFactory.get(prov_name, API_KEYS)
    if not prov or not await prov.health_check():
        return None, None, []

    log.info(f"🤖 Tool calling: {prov_name}/{model}")
    resp = await prov.chat(messages, model, tools=TOOLS_LIST, tool_choice="auto")

    if not resp.success:
        return None, None, []

    tool_calls = getattr(resp, "tool_calls", None)
    if not tool_calls:
        return resp, None, []

    max_rounds = 3
    current_messages = list(messages)
    tools_used = []
    pending_actions = []

    for round_num in range(max_rounds):
        current_messages.append({
            "role": "assistant",
            "content": resp.content or "",
            "tool_calls": tool_calls
        })

        for tc in tool_calls:
            fn_name = tc.get("function", {}).get("name", "")
            fn_args_str = tc.get("function", {}).get("arguments", "{}")
            tool_call_id = tc.get("id", f"call_{round_num}")

            try:
                fn_args = json.loads(fn_args_str)
            except (json.JSONDecodeError, TypeError):
                fn_args = {"query": fn_args_str}

            # Capture music action manually
            if fn_name == "play_music":
                pending_actions.append({
                    "type": "music",
                    "action": fn_args.get("action", "play"),
                    "query": fn_args.get("query", ""),
                })

            # Inject server info for get_server_info tool
            if fn_name == "get_server_info" and settings:
                fn_args["_server_info"] = settings.get("server_info", {})
            tool_result = await execute_tool_call(fn_name, fn_args)

            # Updated action parser
            try:
                result_data = json.loads(tool_result)
                if isinstance(result_data, dict):
                    action_type = result_data.get("type")
                    if action_type in ("download", "image", "upload_file", "reminder", "send_message", "get_server_info"):
                        pending_actions.append(result_data)
            except (json.JSONDecodeError, TypeError):
                pass

            log.info(f"✅ Round {round_num + 1}: {fn_name}")
            tools_used.append(fn_name)

            current_messages.append({
                "role": "tool",
                "tool_call_id": tool_call_id,
                "content": tool_result if not tool_result.startswith("{") else f"Tool result: {tool_result}"
            })

        # Round selesai, minta AI respond
        resp = await prov.chat(current_messages, model)

        if not resp.success:
            if "tool" in str(resp.error).lower():
                log.warning("Tool error, retrying without tool context...")
                clean_messages = [m for m in current_messages if m.get("role") != "tool" and "tool_calls" not in m]
                clean_messages.append({"role": "user", "content": "Based on the information gathered, please provide your response."})
                resp = await prov.chat(clean_messages, model)
            if not resp.success:
                return None, None, pending_actions

        tool_calls = getattr(resp, "tool_calls", None)
        if not tool_calls:
            _log_request(guild_id, prov_name, model, True, resp.latency)
            tool_icons = {
                "web_search": "🔍", "get_time": "🕐", "get_weather": "🌤️",
                "calculate": "🔢", "translate": "🌐", "play_music": "🎵",
                "fetch_url": "📄", "generate_image": "🖼️"
            }
            unique_tools = list(dict.fromkeys(tools_used))
            icons = "".join(tool_icons.get(t, "🔧") for t in unique_tools)
            note = f"{icons} Auto-tools via {prov_name}/{model}" if tools_used else None
            return resp, note, pending_actions

    _log_request(guild_id, prov_name, model, True, resp.latency)
    return resp, f"🔧 Auto-tools ({max_rounds} rounds) via {prov_name}/{model}", pending_actions


# ============================================================
# FALLBACK
# ============================================================

async def execute_with_fallback(messages, mode, preferred_provider, preferred_model, guild_id=0):
    chain = [(preferred_provider, preferred_model)]
    for item in FALLBACK_CHAINS.get(mode, FALLBACK_CHAINS["normal"]):
        if item[0] not in ["duckduckgo", "tavily", "brave", "serper", "jina"] and item not in chain:
            chain.append(item)
    fallback_note, orig_p, orig_m, is_fb = None, preferred_provider, preferred_model, False
    for pname, mid in chain:
        prov = ProviderFactory.get(pname, API_KEYS)
        if not prov or not await prov.health_check(): continue
        log.info(f"Trying {pname}/{mid}")
        resp = await prov.chat(messages, mid)
        if resp.success:
            _log_request(guild_id, pname, mid, True, resp.latency, is_fb)
            if is_fb: fallback_note = f"⚡ {orig_p}/{orig_m} → {pname}/{mid}"
            return resp, fallback_note
        log.warning(f"Failed: {pname}/{mid}")
        _log_request(guild_id, pname, mid, False, resp.latency, is_fb, resp.error)
        is_fb = True
    return AIResponse(False, "Semua provider tidak tersedia.", "none", "none", error="exhausted"), None

# ============================================================
# SYSTEM PROMPTS
# ============================================================

SYSTEM_PROMPTS = {
    "normal": """You are a helpful AI assistant in a Discord server.
You can see who is talking by their name in [brackets].
Multiple users may be chatting — address them by name when appropriate.
Remember full conversation context. Respond in user's language. Be concise and friendly.

When you receive tool results, use that information naturally in your answer.
Do NOT list sources, URLs, or citations. Do NOT use numbered references.
Just answer naturally as if you already knew the information.
Never say "I cannot access real-time data" when tool results are provided.

IMPORTANT — URL HANDLING:
When a user shares ANY URL/link (https://...), you MUST use the fetch_url tool to read it.
NEVER say "I cannot access the link" — always try fetch_url first.

IMPORTANT — FILE HANDLING:
When a user uploads a file, the file content is included in their message.
Read the content and help them with whatever they ask (explain, review, fix, etc).
You can also CREATE files: Word (.docx), Excel (.xlsx), code files (.py, .js, etc).

IMPORTANT — EXCEL FORMAT:
When creating Excel files, you MUST provide content as JSON array of arrays.
First row = headers. Example: [["Name","Age"],["John",25],["Jane",30]]

IMPORTANT — WORD FORMAT:
Use markdown formatting: # Heading, ## Subheading, - bullets, 1. numbered lists.

For music: you can play, skip, pause, resume, and stop music.
If user is NOT in a voice channel, tell them to join one first.

For reminders: you can set one-time or recurring daily reminders.""",

    "reasoning": """You are a reasoning AI. Think step by step.
Multiple users may ask questions — keep track of who asked what.
Do not use <think> tags. Explain naturally. Respond in user's language.""",

    "search": """You are a helpful AI assistant with access to current information.
Answer the user's question naturally using the search results provided.
Do NOT list sources, URLs, or citations. Do NOT use numbered references.
Just incorporate the information naturally into your response.
Respond in user's language. Be conversational and helpful.""",

    "with_skill": """You are a helpful AI assistant.
Present tool results naturally as part of your response.
Do NOT list sources or references. Just answer naturally.
Respond in the same language as the user.""",
}

# ============================================================
# VISION — Process images with AI (Dynamic from settings)
# ============================================================

async def process_with_vision(content: str, image_urls: list, settings: dict) -> Optional[str]:
    """Process message with image using the model from settings"""
    
    # Get provider & model from user settings
    mode = settings.get("active_mode", "normal")
    profile = settings.get("profiles", {}).get(mode, {})
    
    prov_name = profile.get("provider", "pollinations")
    model_id = profile.get("model", "openai")
    
    log.info(f"👁️ Vision using settings: {prov_name}/{model_id}")
    
    # Build message with images (OpenAI vision format)
    user_content = [{"type": "text", "text": content or "Describe this image / Jelaskan gambar ini"}]
    
    for url in image_urls[:3]:  # Max 3 images
        user_content.append({
            "type": "image_url",
            "image_url": {"url": url}
        })
    
    messages = [
        {
            "role": "system", 
            "content": "You are a helpful AI assistant that can see and analyze images. Respond in the same language as the user. Be detailed but concise."
        },
        {"role": "user", "content": user_content}
    ]
    
    # Try with user's selected provider/model first
    prov = ProviderFactory.get(prov_name, API_KEYS)
    if prov:
        try:
            if await prov.health_check():
                log.info(f"👁️ Trying vision: {prov_name}/{model_id}")
                resp = await prov.chat(messages, model_id)
                
                if resp.success and resp.content:
                    log.info(f"👁️ Vision success via {prov_name}/{model_id}")
                    return resp.content
                else:
                    log.warning(f"👁️ Vision failed {prov_name}/{model_id}: {resp.error}")
        except Exception as e:
            log.warning(f"👁️ Vision error {prov_name}: {e}")
    
    # Fallback: try other vision-capable providers
    VISION_FALLBACKS = [
        ("pollinations", "openai"),        # GPT-5 Mini - supports vision
        ("pollinations", "gemini"),        # Gemini 2.5 Pro
        ("gemini", "gemini-2.0-flash"),    # Direct Gemini
        ("groq", "llama-4-scout-17b-16e-instruct"),  # Llama 4 Scout
    ]
    
    for fb_prov, fb_model in VISION_FALLBACKS:
        # Skip if same as already tried
        if fb_prov == prov_name and fb_model == model_id:
            continue
        
        fb_provider = ProviderFactory.get(fb_prov, API_KEYS)
        if not fb_provider:
            continue
        
        try:
            if not await fb_provider.health_check():
                continue
            
            log.info(f"👁️ Fallback vision: {fb_prov}/{fb_model}")
            resp = await fb_provider.chat(messages, fb_model)
            
            if resp.success and resp.content:
                log.info(f"👁️ Vision success via fallback {fb_prov}/{fb_model}")
                return resp.content
        except Exception as e:
            log.warning(f"👁️ Fallback error {fb_prov}: {e}")
            continue
    
    return None


# ============================================================
# MAIN HANDLER
# ============================================================

async def handle_message(content: str, settings: Dict, channel_id: int = 0, user_id: int = 0, user_name: str = "User") -> Dict:
    mode = settings.get("active_mode", "normal")
    guild_id = settings.get("guild_id", 0)

    history = get_conversation(guild_id, channel_id, limit=30)

        # =========================================================
    # STEP 0: Read file attachments (if any)
    # =========================================================
    
    file_context = ""
    image_urls = []
    attachments_data = settings.get("attachments", [])
    
    if attachments_data:
        for att in attachments_data[:3]:  # Max 3 files
            file_content = await read_uploaded_file(att["url"], att["filename"])
            if file_content:
                # Check if it's an image
                try:
                    parsed = json.loads(file_content)
                    if parsed.get("type") == "image_attachment":
                        image_urls.append(parsed["url"])
                        file_context += f"\n\n[Image attached: {parsed['filename']}]"
                        continue
                except (json.JSONDecodeError, TypeError):
                    pass
                
                file_context += f"\n\n{file_content}"
        
        if file_context:
            content = content + file_context
            log.info(f"📎 Attached {len(attachments_data)} file(s) to message")
    
    # Store image URLs for vision processing
    settings["image_urls"] = image_urls

    # =========================================================
    # STEP 0B: Process images with Vision AI
    # =========================================================
    
    image_urls = settings.get("image_urls", [])
    if image_urls:
        log.info(f"👁️ Processing {len(image_urls)} image(s) with vision AI")
        
        vision_result = await process_with_vision(content, image_urls, settings)
        
        if vision_result:
            save_message(guild_id, channel_id, user_id, user_name, "user", f"{content} [+{len(image_urls)} image(s)]")
            save_message(guild_id, channel_id, user_id, user_name, "assistant", vision_result)
            return {"text": vision_result, "fallback_note": "👁️ Vision AI", "actions": []}
        else:
            return {
                "text": "Maaf, saya tidak bisa memproses gambar saat ini. Coba lagi nanti.",
                "fallback_note": None,
                "actions": []
            }

    # =========================================================
    # STEP 1: Try smart skills (time, weather, calendar)
    # =========================================================

    skill_result = None
    try:
        from skills.detector import SkillDetector
        skill_result = await SkillDetector.detect_and_execute(content)
    except Exception as e:
        log.warning(f"Skill detection error: {e}")

    if skill_result:
        profile = settings.get("profiles", {}).get(mode, {"provider": "groq", "model": "llama-3.3-70b-versatile"})
        prov, mid = profile.get("provider", "groq"), profile.get("model", "llama-3.3-70b-versatile")

        msgs = [
            {"role": "system", "content": SYSTEM_PROMPTS["with_skill"]},
            *[{"role": m["role"], "content": m["content"]} for m in history],
            {"role": "user", "content": f"[{user_name}] bertanya: {content}\n\nHasil tool:\n{skill_result}\n\nSampaikan informasi ini secara natural."}
        ]

        resp, fb_note = await execute_with_fallback(msgs, mode, prov, mid, guild_id)
        text = strip_think_tags(resp.content) if resp.success else skill_result

        save_message(guild_id, channel_id, user_id, user_name, "user", content)
        save_message(guild_id, channel_id, user_id, user_name, "assistant", text)

        return {"text": text, "fallback_note": fb_note if resp.success else None, "actions": []}

    # =========================================================
    # STEP 2: Auto-detect mode
    # =========================================================

    if settings.get("auto_detect"):
        detected = ModeDetector.detect(content)
        if detected != "normal":
            mode = detected

    # =========================================================
    # STEP 2B: Auto Tool Calling
    # =========================================================

    profile = settings.get("profiles", {}).get(mode, {"provider": "groq", "model": "llama-3.3-70b-versatile"})
    prov, mid = profile.get("provider", "groq"), profile.get("model", "llama-3.3-70b-versatile")

    from core.providers import supports_tool_calling
    if supports_tool_calling(prov):
        system_prompt = SYSTEM_PROMPTS.get(mode, SYSTEM_PROMPTS["normal"])
        formatted_history = []
        for msg in history:
            if msg["role"] == "user" and msg.get("user_name"):
                formatted_history.append({"role": "user", "content": f"[{msg['user_name']}]: {msg['content']}"})
            else:
                formatted_history.append({"role": msg["role"], "content": msg["content"]})

        voice_ctx = ""
        if settings.get("user_in_voice"):
            voice_ctx = f" [in voice channel: {settings.get('user_voice_channel', 'yes')}]"
        else:
            voice_ctx = " [not in voice channel]"

        # ── URL Detection: inject hint so AI uses fetch_url ──
        user_content = f"[{user_name}]{voice_ctx}: {content}"
        urls = re.findall(r'https?://[^\s<>"\']+', content)
        if urls:
            url_hint = f"\n\n[System hint: User shared {len(urls)} URL(s). Use fetch_url tool to read the content. URLs: {', '.join(urls)}]"
            user_content += url_hint

        tool_msgs = [
            {"role": "system", "content": system_prompt},
            *formatted_history,
            {"role": "user", "content": user_content}
        ]

        tool_resp, tool_note, tool_actions = await handle_with_tools(tool_msgs, prov, mid, guild_id, settings)
        if tool_resp and tool_resp.success:
            text = strip_think_tags(tool_resp.content) or "Tidak ada jawaban."
            save_message(guild_id, channel_id, user_id, user_name, "user", content)
            save_message(guild_id, channel_id, user_id, user_name, "assistant", text)
            return {"text": text, "fallback_note": tool_note, "actions": tool_actions}

    # =========================================================
    # STEP 3: Regular AI chat (Fallback)
    # =========================================================

    profile = settings.get("profiles", {}).get(mode, {"provider": "groq", "model": "llama-3.3-70b-versatile"})
    prov, mid = profile.get("provider", "groq"), profile.get("model", "llama-3.3-70b-versatile")
    system_prompt = SYSTEM_PROMPTS.get(mode, SYSTEM_PROMPTS["normal"])

    formatted_history = []
    for msg in history:
        if msg["role"] == "user" and msg.get("user_name"):
            formatted_history.append({"role": "user", "content": f"[{msg['user_name']}]: {msg['content']}"})
        else:
            formatted_history.append({"role": msg["role"], "content": msg["content"]})

    if mode == "search" and not is_grounding_model(prov, mid):
        search_res = await do_search(content, profile.get("engine", "duckduckgo"))
        msgs = [
            {"role": "system", "content": system_prompt},
            *formatted_history,
            {"role": "user", "content": f"[{user_name}]: {content}\n\nHasil pencarian:\n{search_res}"}
        ]
    else:
        msgs = [
            {"role": "system", "content": system_prompt},
            *formatted_history,
            {"role": "user", "content": f"[{user_name}]: {content}"}
        ]

    resp, fb_note = await execute_with_fallback(msgs, mode, prov, mid, guild_id)

    if resp.success:
        text = strip_think_tags(resp.content) or "Tidak ada jawaban."
        save_message(guild_id, channel_id, user_id, user_name, "user", content)
        save_message(guild_id, channel_id, user_id, user_name, "assistant", text)
        return {"text": text, "fallback_note": fb_note, "actions": []}
    return {"text": resp.content, "fallback_note": None, "actions": []}
